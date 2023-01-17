from docx import Document
from openpyxl.workbook import Workbook
from openpyxl import load_workbook
from country_list import *


c_tries = dict(countries_for_language('en'))

the_countries = ["Bahamas",
                "Cayman Islands",
                "Central African Republic",
                "Channel Islands",
                "Comoros",
                "Czech Republic",
                "Dominican Republic",
                "Falkland Islands",
                "Gambia",
                "Isle of Man",
                "Ivory Coast",
                "Leeward Islands",
                "Maldives",
                "Marshall Islands",
                "Netherlands",
                "Netherlands Antilles",
                "Philippines",
                "Solomon Islands",
                "Turks & Caicos Islands",
                "United Arab Emirates",
                "United Kingdom",
                "United States",
                "Virgin Islands"]

countries = []

for country in c_tries:
    if c_tries[country] in the_countries:
        countries.append(f"the {c_tries[country]}")
    else:
        countries.append(c_tries[country])



document = Document("_ctp1.docx")
document1 = Document("_ctp.docx")
# CHECK if time is in the template --> put on screen the necessary time part

templates = []
subjects = []   
bodytexts = []

curr_paragraph = ""
x = 0

for paragraph in document.paragraphs:
    if paragraph.style.name == "Heading 1":
        templates.append(paragraph.text)
        x = 1
    elif paragraph.style.name == "Heading 2":
        subjects.append(paragraph.text)
        x = 2
    elif paragraph.style.name == "Normal":
        x = 0

    if x == 0:
        curr_paragraph = curr_paragraph + paragraph.text
    elif x == 1 and curr_paragraph != "":
        bodytexts.append(curr_paragraph)
        curr_paragraph = ""

templates1 = []
subjects1 = []
bodytexts1 = []
curr_paragraph = ""
x = 0

for paragraph in document1.paragraphs:
    if paragraph.style.name == "Heading 1":
        templates1.append(paragraph.text)
        x = 1  
    elif paragraph.style.name == "Heading 2":
        subjects1.append(paragraph.text)
        x = 2
    elif paragraph.style.name == "Normal":
        x = 0

    if x == 0:
        curr_paragraph = curr_paragraph + paragraph.text
    elif x == 1:
        bodytexts1.append(curr_paragraph)
        curr_paragraph = ""

# get the teams from the excel

# wb = load_workbook("_GES.xlsx")
wb = load_workbook("GES_clientallocation.xlsx")
sheets = wb.sheetnames
ca = wb[sheets[0]]
email = wb[sheets[1]]

eng_name = []
engagement = []
manager = []
reviewer = []
preparer = []
comp_spec = []
emails = []
engagements = []

i=2
while ca[f"C{i}"].value != None:
    eng_name.append(ca[f"C{i}"].value)
    manager.append(ca[f"K{i}"].value)
    reviewer.append(ca[f"J{i}"].value)
    preparer.append(ca[f"I{i}"].value)
    comp_spec.append(ca[f"H{i}"].value)
    i += 1

i=2
while email[f"A{i}"].value != None:
    name = email[f"A{i}"].value
    address = email[f"D{i}"].value
    emails.append({"name":name,"email":address})
    i += 1

engagement_teams = []

for i in range(len(eng_name)):
    man = ""
    rev = ""
    prep = ""
    comp = ""
    for email in emails:
        if comp_spec[i] == email["name"]:
            comp = email
        if preparer[i] == email["name"]:
            prep = email
        if reviewer[i] == email["name"]:
            rev = email
        if manager[i] == email["name"]:
            man = email
    team = {"manager":man,"reviewer":rev,"preparer":prep, "compspec":comp}
    eng = {"engagement":eng_name[i],"team":team}
    engagement_teams.append(eng)